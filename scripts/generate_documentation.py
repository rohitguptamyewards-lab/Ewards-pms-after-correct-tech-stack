from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
DOCS_DIR.mkdir(exist_ok=True)

MARKDOWN_PATH = DOCS_DIR / "EWards_PMS_Detailed_Documentation.md"
DOCX_PATH = DOCS_DIR / "EWards_PMS_Detailed_Documentation.docx"


WEB_ROUTES: list[tuple[str, str, str, str]] = [
    ("GET", "/login", "LoginController@showLoginForm", "Login page rendered with Inertia."),
    ("POST", "/login", "LoginController@authenticate", "Session-based login using TeamMember credentials."),
    ("POST", "/logout", "LoginController@logout", "Session invalidation and logout."),
    ("GET", "/", "DashboardController@index", "Role-based dashboard entry point."),
    ("GET", "/projects", "ProjectController@index", "Project list view."),
    ("GET", "/projects/custom-worklog", "ProjectController@customWorklogIndex", "Custom work project listing."),
    ("GET", "/projects/create", "ProjectController@create", "Project creation screen."),
    ("POST", "/projects", "ProjectController@store", "Create new project."),
    ("GET", "/projects/board", "ProjectController@board", "Kanban-style board view."),
    ("GET", "/projects/{id}", "ProjectController@show", "Project detail workspace."),
    ("PUT", "/projects/{id}", "ProjectController@update", "Update project fields."),
    ("GET", "/projects/{id}/replicate", "ProjectController@replicate", "Open create page pre-filled from an existing project."),
    ("GET", "/work-logs", "WorkLogController@index", "Work log list."),
    ("GET", "/work-logs/create", "WorkLogController@create", "Work log creation screen."),
    ("POST", "/work-logs", "WorkLogController@store", "Create work log."),
    ("GET", "/reports/dashboard", "ReportController@dashboard", "Comprehensive reporting dashboard."),
    ("GET", "/reports/projects", "ReportController@projects", "Project reporting screen."),
    ("GET", "/reports/workers", "ReportController@workers", "Worker reporting screen."),
    ("GET", "/automations", "AutomationController@index", "Automation management screen."),
    ("GET", "/release-notes", "ReleaseNoteController@allIndex", "All release notes page."),
    ("GET", "/projects/{projectId}/release-notes", "ReleaseNoteController@index", "Project-scoped release notes page."),
    ("GET", "/team-members", "TeamMemberController@index", "Team member listing."),
    ("GET", "/team-members/create", "TeamMemberController@create", "Create team member screen."),
    ("POST", "/team-members", "TeamMemberController@store", "Create team member record."),
]

API_ROUTES: list[tuple[str, str, str, str]] = [
    ("GET", "/api/v1/projects", "ProjectController@index", "Project list JSON."),
    ("POST", "/api/v1/projects", "ProjectController@store", "Create project JSON."),
    ("GET", "/api/v1/projects/{id}", "ProjectController@show", "Project detail JSON."),
    ("PUT", "/api/v1/projects/{id}", "ProjectController@update", "Update project JSON."),
    ("GET", "/api/v1/projects/{id}/children", "ProjectController@children", "Load subtasks."),
    ("POST", "/api/v1/projects/{id}/attachments", "ProjectController@uploadAttachment", "Upload project attachment."),
    ("DELETE", "/api/v1/attachments/{id}", "ProjectController@deleteAttachment", "Delete project attachment."),
    ("GET", "/api/v1/projects/{projectId}/planners", "PlannerController@index", "Planner items by project."),
    ("POST", "/api/v1/planners", "PlannerController@store", "Create planner item."),
    ("PUT", "/api/v1/planners/{id}", "PlannerController@update", "Update planner item."),
    ("DELETE", "/api/v1/planners/{id}", "PlannerController@destroy", "Delete planner item."),
    ("POST", "/api/v1/projects/{projectId}/planners/reorder", "PlannerController@reorder", "Reorder planner items."),
    ("GET", "/api/v1/projects/{projectId}/workers", "WorkerController@index", "List project workers."),
    ("POST", "/api/v1/projects/{projectId}/workers/owner", "WorkerController@assignOwner", "Assign or change owner."),
    ("POST", "/api/v1/projects/{projectId}/workers/contributor", "WorkerController@addContributor", "Add contributor."),
    ("DELETE", "/api/v1/projects/{projectId}/workers/{userId}", "WorkerController@removeContributor", "Remove contributor."),
    ("GET", "/api/v1/projects/{projectId}/stage", "StageController@show", "Get current stage."),
    ("PUT", "/api/v1/projects/{projectId}/stage", "StageController@update", "Add new stage entry."),
    ("GET", "/api/v1/projects/{projectId}/stage/history", "StageController@history", "Get stage history."),
    ("GET", "/api/v1/projects/{projectId}/updates", "UpdateController@index", "Get comments / updates."),
    ("POST", "/api/v1/projects/{projectId}/updates", "UpdateController@store", "Create comment / update."),
    ("GET", "/api/v1/projects/{projectId}/tickets", "TicketLinkController@index", "Get linked tickets."),
    ("POST", "/api/v1/projects/{projectId}/tickets", "TicketLinkController@store", "Create linked ticket."),
    ("DELETE", "/api/v1/tickets/{id}", "TicketLinkController@destroy", "Delete linked ticket."),
    ("GET", "/api/v1/projects/{projectId}/transfers", "TransferController@index", "Get transfer history."),
    ("POST", "/api/v1/projects/{projectId}/transfers", "TransferController@store", "Transfer ownership."),
    ("GET", "/api/v1/projects/{projectId}/blockers", "BlockerController@index", "Get blockers."),
    ("POST", "/api/v1/projects/{projectId}/blockers", "BlockerController@store", "Create blocker."),
    ("PUT", "/api/v1/blockers/{id}/resolve", "BlockerController@resolve", "Resolve blocker."),
    ("GET", "/api/v1/dashboard/activity-report", "DashboardController@activityReport", "Team activity report JSON."),
    ("GET", "/api/v1/reports/projects", "ReportController@projects", "Projects report JSON."),
    ("GET", "/api/v1/reports/workers", "ReportController@workers", "Workers report JSON."),
    ("GET", "/api/v1/reports/dashboard", "ReportController@dashboard", "Report dashboard JSON."),
    ("GET", "/api/v1/reports/projects/{projectId}/worklogs", "ReportController@projectWorklogs", "Project worklog analytics."),
    ("GET", "/api/v1/reports/members/{memberId}/worklogs", "ReportController@memberWorklogs", "Member worklog analytics."),
    ("GET", "/api/v1/work-logs", "WorkLogController@index", "Work log list JSON."),
    ("POST", "/api/v1/work-logs", "WorkLogController@store", "Create work log JSON."),
    ("PUT", "/api/v1/work-logs/{id}", "WorkLogController@update", "Update work log JSON."),
    ("DELETE", "/api/v1/work-logs/{id}", "WorkLogController@destroy", "Delete work log JSON."),
    ("PUT", "/api/v1/team-members/{id}", "TeamMemberController@update", "Update team member."),
    ("POST", "/api/v1/team-members/{id}/toggle-active", "TeamMemberController@toggleActive", "Toggle active status."),
    ("GET", "/api/v1/automations", "AutomationController@index", "Automation list."),
    ("POST", "/api/v1/automations", "AutomationController@store", "Create automation."),
    ("PUT", "/api/v1/automations/{id}", "AutomationController@update", "Update automation."),
    ("DELETE", "/api/v1/automations/{id}", "AutomationController@destroy", "Delete automation."),
    ("POST", "/api/v1/automations/{id}/toggle", "AutomationController@toggle", "Enable/disable automation."),
    ("POST", "/api/v1/automations/{id}/run", "AutomationController@runNow", "Force-run scheduled automation."),
    ("GET", "/api/v1/automations/{id}/logs", "AutomationController@logs", "Automation log history."),
    ("POST", "/api/v1/release-notes/setup", "ReleaseNoteController@setupTables", "Create release note tables if missing."),
    ("GET", "/api/v1/projects/{projectId}/release-notes", "ReleaseNoteController@index", "List project release notes."),
    ("POST", "/api/v1/projects/{projectId}/release-notes", "ReleaseNoteController@store", "Create release note."),
    ("PUT", "/api/v1/release-notes/{id}", "ReleaseNoteController@update", "Update release note."),
    ("DELETE", "/api/v1/release-notes/{id}", "ReleaseNoteController@destroy", "Delete release note."),
    ("POST", "/api/v1/release-notes/{id}/files", "ReleaseNoteController@uploadFiles", "Upload release note files."),
    ("DELETE", "/api/v1/release-note-files/{fileId}", "ReleaseNoteController@deleteFile", "Delete release note file."),
    ("POST", "/api/v1/release-notes/{id}/links", "ReleaseNoteController@addLink", "Add release note link."),
    ("DELETE", "/api/v1/release-note-links/{linkId}", "ReleaseNoteController@deleteLink", "Delete release note link."),
    ("GET", "/api/v1/notifications", "NotificationController@index", "Fetch current user's notifications."),
    ("PUT", "/api/v1/notifications/{id}/read", "NotificationController@markRead", "Mark single notification read."),
    ("PUT", "/api/v1/notifications/read-all", "NotificationController@markAllRead", "Mark all notifications read."),
    ("GET", "/api/v1/team-members", "Closure", "Dropdown data of active team members."),
]

ROLE_SUMMARIES: dict[str, list[str]] = {
    "Manager": [
        "Highest operational role in the implemented system.",
        "Can access manager dashboard, team members, reports, workers report, projects report, automations, release note deletion, team activity report, ownership assignment, and ownership transfer.",
        "Can view all work logs.",
        "Can create team members and toggle team member active status.",
    ],
    "Analyst Head": [
        "Admin-like role with nearly the same reach as manager in the current implementation.",
        "Can access manager dashboard, team member management, projects report, workers report, report dashboard, automations, release note deletion, ownership assignment, and ownership transfer.",
        "Can create projects and view team member listing.",
        "Cannot use the manager-only JSON placeholder route `/api/v1/dashboard/manager`, but can use the manager web dashboard.",
    ],
    "Analyst": [
        "Lead functional role focused on project creation and management.",
        "Can access the manager-style dashboard, project creation, project replication, projects report, report dashboard, team member listing, and ownership transfer.",
        "Cannot manage automations, cannot create team members, cannot assign owner/contributor through WorkerController, and cannot view the restricted team activity report.",
        "Can only see their own work logs in the work log module because privileged work-log access excludes analysts.",
    ],
    "Senior Developer": [
        "Technical leadership role with enhanced reporting and visibility.",
        "Uses the employee dashboard path, but gets extra access to workers report, report dashboard sensitive sections, and team activity report.",
        "Can view all work logs.",
        "Does not get project creation or automation management rights through the implemented backend.",
    ],
    "Developer": [
        "Standard delivery role for assigned work.",
        "Gets the employee dashboard and base navigation.",
        "Can create work logs, access assigned projects through filtered list pages, and receive notifications tied to stage changes and assignments.",
        "Cannot create projects, cannot manage team members, cannot manage automations, and cannot view all work logs.",
    ],
    "Employee": [
        "Lowest permission role kept partly for backward compatibility.",
        "Gets the employee dashboard and base navigation.",
        "Can log work and view base modules available in the sidebar.",
        "Explicitly blocked from `ProjectController@update`, which means employees cannot edit project details.",
    ],
}

ACCESS_MATRIX: dict[str, str] = {
    "Dashboard": "All authenticated users. Manager-style dashboard for manager/analyst_head/analyst; employee-style dashboard for others.",
    "Projects list / board / custom work": "All authenticated users can open the module. Manager, analyst_head, analyst see global lists; all other roles see only assigned / related records in index endpoints.",
    "Project creation / replication": "Manager, analyst_head, analyst only.",
    "Project update": "All authenticated users except employee. No owner-only restriction is enforced in the controller.",
    "Project detail view": "Any authenticated user who knows the route can load it; there is no per-project authorization gate in `show()`.",
    "Planner CRUD": "Any authenticated user by backend route protection only. No role gate exists in PlannerController.",
    "Worker ownership / contributor management": "Manager and analyst_head only.",
    "Ownership transfer": "Manager, analyst_head, analyst.",
    "Stage changes": "Any authenticated user by backend route protection only. No role gate exists in StageController.",
    "Project updates / comments": "Any authenticated user by backend route protection only.",
    "Ticket links": "Any authenticated user by backend route protection only.",
    "Blockers": "Any authenticated user by backend route protection only.",
    "Work log create / update / delete": "All authenticated users for their own records. Manager, analyst_head, senior_developer can also view and modify others' logs.",
    "Team member list": "Manager, analyst_head, analyst.",
    "Team member create / update / toggle": "Manager and analyst_head only.",
    "Projects report": "Manager, analyst_head, analyst.",
    "Workers report": "Any authenticated user can hit the endpoint, but the UI exposes it mainly to manager, analyst_head, and senior_developer.",
    "Report dashboard": "Manager, analyst_head, analyst, senior_developer. Sensitive sections are only visible to manager, analyst_head, senior_developer.",
    "Team activity report": "Manager, analyst_head, senior_developer only.",
    "Release note create / update": "Any authenticated user can create and update release notes because `canCreate()` only checks `auth()->check()`.",
    "Release note delete / setup": "Manager and analyst_head only.",
    "Automations": "Manager and analyst_head only.",
    "Notifications": "Current authenticated user only.",
}

VALIDATION_RULES: dict[str, list[str]] = {
    "Project create": [
        "Required: `name`, `owner_id`.",
        "Optional structured fields: `description`, `objective`, `tags`, `status`, `priority`, `work_type`, `task_type`, `custom_task_type`, `ticket_link`, `document_link`, `ai_chat_link`, `parent_id`, `start_date`, `due_date`, `linked_project_ids`.",
        "User assignments must exist in `team_members`: `analyst_id`, `analyst_testing_id`, `developer_id`, `owner_id`.",
        "If `due_date` is present it must be after or equal to `start_date`.",
        "Maximum nesting depth is enforced in the controller: top level depth 0, overall maximum 4 levels.",
    ],
    "Project update": [
        "Same field family as create, but all fields are optional.",
        "Employees are blocked before validation in the controller.",
    ],
    "Planner create": [
        "Required: `project_id`, `title`.",
        "Optional: `description`, `milestone_flag`, `assigned_to`, `due_date`, `status`.",
        "If no assignee is supplied, `PlannerService` falls back to the project owner.",
    ],
    "Work log create": [
        "Requires either `project_id` or `project_name`.",
        "Required: `log_date`, `start_time`, `end_time`, `note`.",
        "Optional: `status`, `blocker`.",
        "End time must be after start time.",
        "If a custom `project_name` is used and no project exists, the system auto-creates a worklog-only project.",
    ],
    "Team member create": [
        "Required: `name`, `email`, `password`, `role`.",
        "Optional: `joined_date`.",
        "Email must be unique in `team_members`.",
        "Password minimum length is 6.",
    ],
    "Automation create": [
        "Required: `name`, `trigger_type`, `trigger_config`, `action_type`, `action_config`.",
        "Allowed trigger types: `schedule`, `stage_change`, `status_change`, `blocker_created`.",
        "Allowed action types: `send_email`, `send_notification`.",
    ],
    "Release note create": [
        "Required: `title`.",
        "Optional: `description`, `files[]`, `links[]`.",
        "File upload limit is 20 MB per file.",
        "Each link row requires a `url`; `label` is optional.",
    ],
}

DATA_MODEL_NOTES: dict[str, str] = {
    "team_members": "Authentication entity. Stores name, email, password hash, role, active flag, joined date, remember token, timestamps, and soft delete state.",
    "sessions": "Session driver backing table used by Laravel session authentication.",
    "projects": "Primary project record with status, priority, owner, creator, tags, linked projects, work type, task type, optional analyst/developer/testing assignments, document links, AI chat link, custom task type flag, parent-child nesting, dates, and soft deletes.",
    "project_planners": "Task / milestone records linked to a project, with assignee, due date, status, and order index.",
    "project_workers": "People assigned to a project, tracked as `owner` or `contributor`, including assignment metadata.",
    "project_stages": "Immutable history of stage changes per project.",
    "project_updates": "Comment / update feed for a project, including system-generated stage change messages.",
    "project_ticket_links": "External or internal ticket identifiers tied to a project.",
    "project_transfers": "Ownership transfer history from one user to another with notes.",
    "project_blockers": "Blockers raised against a project, including creator, resolver, and active/resolved status.",
    "work_logs": "Time entries linked to project and user, with note, blocker text, hours spent, optional status, project stage snapshot, and soft deletes.",
    "project_attachments": "Files uploaded directly to a project.",
    "notifications": "In-app notifications addressed to a single user, with type, title, message, optional JSON data, and read timestamp.",
    "automations": "Automation definitions with trigger type/config, action type/config, active flag, creator, and last run timestamp.",
    "automation_logs": "Execution history of automations, including success/failure status and JSON details.",
    "release_notes": "Release note header record linked to a project and creator.",
    "release_note_files": "Files uploaded against a release note.",
    "release_note_links": "External URLs associated with a release note.",
    "audit_logs": "Generic audit trail written by the `Auditable` trait for audited models.",
    "cache / cache_locks": "Laravel database cache backend tables.",
}

FUNCTION_PURPOSES: dict[str, list[str]] = {
    "Authentication": [
        "Session login handled by `LoginController` with TeamMember as the auth model.",
        "No password reset backend exists; the login page has a placeholder `Forgot Password?` link only.",
    ],
    "Dashboard": [
        "`DashboardController@index` routes users to manager or employee dashboard layouts based on role.",
        "Manager dashboard focuses on project counts, blockers, overdue planner items, recent projects, and optionally the team activity report.",
        "Employee dashboard focuses on assigned projects, pending planner items, blockers created by the user, and optional team activity reporting for elevated technical roles.",
    ],
    "Projects": [
        "Project list supports filters for status, priority, search, owner, and work type.",
        "Board view groups projects into workflow columns derived from stage families.",
        "Custom Work is a dedicated view for projects created automatically from work logs.",
        "Project create supports replication, dependency links, nested subtasks, dates, role assignments, and external links.",
        "Project detail aggregates planners, workers, updates, blockers, tickets, attachments, transfers, stage history, report data, and release notes.",
        "Assignment emails are sent when analyst, testing analyst, or developer assignments are created or changed.",
    ],
    "Planner": [
        "Planner items act as ordered project tasks with optional milestone flag, assignee, status, and due date.",
        "Reordering is persisted by explicit order index updates.",
    ],
    "Workers and Ownership": [
        "Project workers can be owners or contributors.",
        "Owner reassignment demotes the previous owner to contributor, updates the project owner, and stamps assignment metadata.",
        "Ownership transfer logs who transferred from whom to whom and why.",
    ],
    "Stages": [
        "Stage changes are append-only: each update creates a new `project_stages` row rather than overwriting a single status field.",
        "Every stage update also writes a system entry into `project_updates`.",
        "Relevant assignees receive email notifications for testing or development stage changes.",
        "Live stage changes notify managers, analyst heads, and senior developers both by email and in-app notification.",
        "Stage changes trigger automation processing for event-based automations.",
    ],
    "Comments, Tickets, and Blockers": [
        "Project updates provide a lightweight comment stream.",
        "Ticket links store external ticket identifiers and source type.",
        "Blockers track active/resolved state and notify key stakeholders when a blocker is created.",
    ],
    "Work Logs": [
        "Work logs calculate `hours_spent` automatically from start and end time.",
        "The create flow can propose the user's last end time on the current day as the next start time.",
        "The repository stores a project stage snapshot on the work log for historical reporting, excluding worklog-only custom projects.",
        "Deleting work logs is soft delete via `deleted_at` rather than hard delete.",
    ],
    "Team Members": [
        "Team members are the only authenticated principal used by the system.",
        "Create flow sends a welcome email containing the plain-text password chosen at creation time.",
        "List view includes project count and cumulative work-log hours per member.",
    ],
    "Reports": [
        "Projects report summarizes planner completion, blockers, current stage, total hours, contributors, and a coarse progress percentage.",
        "Workers report summarizes active projects and worklog hours by member.",
        "Report dashboard combines timelines, deadlines, project worklog breakdowns, team utilization, and recent activity feeds.",
    ],
    "Release Notes": [
        "Release notes can be viewed across all projects or within a specific project.",
        "If release-note tables are missing, the UI falls back into a migration-pending mode.",
        "Release notes support attached files and named links.",
    ],
    "Notifications": [
        "Notifications are polled from the frontend every 30 seconds through `AppLayout.vue`.",
        "Users can mark a single notification or all notifications as read.",
        "Notification payloads can contain a route link for deep navigation back to the relevant project.",
    ],
    "Automations": [
        "Scheduled automations are executed through the artisan command `automations:process`.",
        "Event-driven automations are triggered from stage changes, project status changes, and blocker creation code paths.",
        "Automations can send emails or in-app notifications to all assignees, specific roles, or specific users.",
        "Runs are logged into `automation_logs` with status and details.",
    ],
    "Audit Logging": [
        "The `Auditable` trait writes created, updated, and deleted events to `audit_logs`.",
        "Auditing is applied to `Project`, `ProjectPlanner`, `TeamMember`, and `WorkLog` models.",
    ],
}

IMPLEMENTATION_CAVEATS: list[str] = [
    "The system installs Laravel Sanctum, but the implemented authentication flow is classic Laravel session auth using the `web` guard and the `team_members` table.",
    "Many API routes are protected only by `auth` middleware and do not apply role checks. Planner CRUD, stage updates, project comments, ticket links, blockers, and project attachments are examples.",
    "Project detail access (`ProjectController@show`) does not verify that the current user is assigned to the project; any authenticated user with the URL can load it.",
    "Project updates are blocked only for the `employee` role. Developers and senior developers can update project details if they can reach the endpoint.",
    "Release note creation and update are available to any authenticated user because `canCreate()` only checks `auth()->check()` even though comments suggest narrower intent.",
    "The login screen advertises forgot password, but there is no reset flow implemented in routes or controllers.",
    "ReportController `workers()` only checks `auth()->check()` even though the sidebar mainly exposes it to manager, analyst_head, and senior_developer.",
]


def collect_php_methods() -> list[tuple[str, list[str]]]:
    inventory: list[tuple[str, list[str]]] = []

    for path in sorted((ROOT / "app").rglob("*.php")):
        text = path.read_text(encoding="utf-8")
        methods = re.findall(r"(?:public|protected|private)\s+function\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(", text)
        if not methods:
            continue
        rel = path.relative_to(ROOT).as_posix()
        inventory.append((rel, methods))

    return inventory


def collect_frontend_pages() -> list[str]:
    pages = []
    for path in sorted((ROOT / "resources/js/Pages").rglob("*.vue")):
        pages.append(path.relative_to(ROOT).as_posix())
    return pages


def collect_layouts_and_components() -> list[str]:
    items = []
    for base in ("resources/js/Layouts", "resources/js/Components"):
        for path in sorted((ROOT / base).rglob("*.vue")):
            items.append(path.relative_to(ROOT).as_posix())
    return items


def collect_enum_values() -> list[tuple[str, list[tuple[str, str]]]]:
    rows = []
    for path in sorted((ROOT / "app/Enums").glob("*.php")):
        text = path.read_text(encoding="utf-8")
        cases = re.findall(r"case\s+([A-Za-z_][A-Za-z0-9_]*)\s*=\s*'([^']+)'", text)
        rows.append((path.stem, cases))
    return rows


def collect_schema_inventory() -> dict[str, list[str]]:
    tables: dict[str, list[str]] = defaultdict(list)

    for path in sorted((ROOT / "database/migrations").glob("*.php")):
        current_table: str | None = None
        for raw_line in path.read_text(encoding="utf-8").splitlines():
            line = raw_line.strip()

            create_match = re.search(r"Schema::create\('([^']+)'", line)
            alter_match = re.search(r"Schema::table\('([^']+)'", line)

            if create_match:
                current_table = create_match.group(1)
                continue

            if alter_match:
                current_table = alter_match.group(1)
                continue

            if current_table and line == "});":
                current_table = None
                continue

            if not current_table or "$table->" not in line:
                continue

            special = re.search(r"\$table->(timestamps|softDeletes|rememberToken)\(", line)
            if special:
                value = special.group(1)
                if value not in tables[current_table]:
                    tables[current_table].append(value)
                continue

            regular = re.search(r"\$table->([A-Za-z_][A-Za-z0-9_]*)\('([^']+)'", line)
            if regular:
                column = regular.group(2)
                if column not in tables[current_table]:
                    tables[current_table].append(column)

    return dict(sorted(tables.items()))


def bullet_block(items: list[str]) -> list[str]:
    return [f"- {item}" for item in items]


def section(title: str, body_lines: list[str]) -> list[str]:
    return [f"## {title}", ""] + body_lines + [""]


def build_markdown() -> str:
    lines: list[str] = []

    lines.extend(
        [
            "# eWards PMS Basic Detailed System Documentation",
            "",
            f"Generated on: {date.today().isoformat()}",
            "",
            "## 1. Executive Summary",
            "",
            "eWards PMS Basic is a Laravel 13 project management and work tracking application using Vue 3, Inertia.js, Tailwind CSS 4, Vite, and MySQL 8. The application is designed as a server-driven SPA where Laravel handles routing, validation, business logic, storage, email, and reporting queries, while Vue pages render the user interface and use Axios for richer interactions against `/api/v1` endpoints.",
            "",
            "The main product areas are authentication, dashboards, projects, planner tasks, worker assignment, stage management, work logs, reporting, team management, release notes, notifications, and automations. The codebase mixes repository/service patterns with direct query-builder controllers, so a complete technical handover needs both module-level business explanation and a raw function inventory. This document covers both.",
            "",
            "## 2. Finalized Tech Stack",
            "",
        ]
    )
    lines.extend(
        bullet_block(
            [
                "Backend Framework: Laravel 13",
                "Frontend Framework: Vue.js 3 plus Inertia.js",
                "UI Layer: Tailwind CSS 4 with custom Vue components",
                "Database: MySQL 8",
                "Build Tooling: Vite 8",
                "HTTP client on frontend: Axios",
                "Email delivery: Brevo mail transport through Symfony mailer",
                "Storage: Laravel public disk for project attachments and release-note files",
                "Authentication: Laravel session auth using TeamMember model, not token-based SPA auth",
                "Deployment shape: Docker-based runtime with PHP CLI app server and Node 20 build step",
            ]
        )
    )
    lines.extend(
        [
            "",
            "## 3. Architecture and Flow",
            "",
        ]
    )
    lines.extend(
        bullet_block(
            [
                "Laravel handles web routes, API routes, validation, auth, business logic, DB access, file storage, emails, and scheduled automation execution.",
                "Inertia bridges Laravel routes to Vue page components so the app behaves like a SPA without a separate frontend API host.",
                "Vue pages under `resources/js/Pages` are mounted by `resources/js/app.js` and use the shared `AppLayout` when authenticated.",
                "Tailwind CSS 4 is configured through `resources/css/app.css` and `vite.config.js`.",
                "The authenticated user and flash messages are injected into every Inertia response by `HandleInertiaRequests` middleware.",
                "Notifications are polled every 30 seconds from the main layout and surfaced in a top-bar dropdown.",
                "Attachments and release-note files are stored on the public disk and exposed through `/storage/...` links.",
                "Automations are stored in database tables and executed either by application events or the `automations:process` artisan command.",
            ]
        )
    )
    lines.extend([""])

    lines.extend(section("4. Roles and Permission Model", []))
    for role_name, details in ROLE_SUMMARIES.items():
        lines.append(f"### {role_name}")
        lines.append("")
        lines.extend(bullet_block(details))
        lines.append("")

    lines.append("### Module Access Matrix")
    lines.append("")
    for module_name, note in ACCESS_MATRIX.items():
        lines.append(f"- {module_name}: {note}")
    lines.append("")

    lines.extend(section("5. User Interface and Navigation", [
        "The authenticated shell is driven by `resources/js/Layouts/AppLayout.vue`. Sidebar visibility is role-sensitive even when backend role checks are sometimes broader or looser.",
        "",
        "Base navigation shown to all authenticated roles:",
    ]))
    lines.extend(bullet_block([
        "Dashboard",
        "Projects",
        "Board",
        "Custom Work",
        "Work Logs",
        "Release Notes",
    ]))
    lines.extend(["", "Additional navigation shown by role:", ""])
    lines.extend(bullet_block([
        "Manager, Analyst Head, Analyst: Team Members, Report Dashboard, Projects Report",
        "Manager, Analyst Head, Senior Developer: Workers Report",
        "Manager, Analyst Head: Automations",
        "Manager, Analyst Head, Senior Developer: Team Activity Report widget inside dashboards",
    ]))
    lines.extend(["", "Interface-wide behavior:", ""])
    lines.extend(bullet_block([
        "Dark mode is persisted in localStorage.",
        "Notifications are polled in the layout and can deep-link back to projects.",
        "Flash success and error messages are rendered globally in the layout.",
        "The login page is a dedicated branded Vue screen, not a Blade-auth scaffold.",
    ]))
    lines.extend([""])

    lines.extend(section("6. Feature Catalogue: What, Why, and How", []))
    for feature_name, details in FUNCTION_PURPOSES.items():
        lines.append(f"### {feature_name}")
        lines.append("")
        lines.extend(bullet_block(details))
        lines.append("")

    lines.append("### Project Module Breakdown")
    lines.append("")
    lines.extend(bullet_block([
        "List view: paginated project grid with filters, quick stage change, quick rename, assignee chips, nested subtasks, and deep links to reports and external links.",
        "Board view: grouped workflow board with high-level stage buckets such as Yet to Start, Documentation, Development, Testing, Live Testing, Ready/Review, and Live.",
        "Create view: supports direct project creation, subtask creation, and project replication with copied data.",
        "Show view: tabbed workspace with Details, Planner, Workers, Report, Attachments, Tickets, Blockers, and Release Notes. The Report tab is inserted conditionally in the frontend for elevated roles.",
        "Custom Work: separate listing for worklog-generated projects marked with `custom_task_type = worklog_custom_project`.",
    ]))
    lines.extend(["", "### Notification and Email Triggers", ""])
    lines.extend(bullet_block([
        "Welcome email: sent when a team member is created.",
        "Project assignment email: sent when analyst, testing analyst, or developer assignment is created or changed.",
        "Stage changed email: sent to testers for testing stages, developers for dev stages, and leadership when a project goes live.",
        "In-app blocker notifications: sent when blockers are created.",
        "In-app stage-change notifications: sent to project participants and leadership when relevant, especially on go-live.",
        "Automation email and in-app notifications: sent when automation trigger criteria are met.",
    ]))
    lines.extend([""])

    lines.extend(section("7. Validation and Business Rules", []))
    for name, rules in VALIDATION_RULES.items():
        lines.append(f"### {name}")
        lines.append("")
        lines.extend(bullet_block(rules))
        lines.append("")

    lines.append("### Additional Business Rules")
    lines.append("")
    lines.extend(bullet_block([
        "Project nesting is limited to 4 total levels.",
        "Owner worker row is auto-created whenever a project is created.",
        "Removing an owner through worker management is blocked; ownership must be transferred first.",
        "Work-log custom projects are excluded from normal project lists, dashboards, and most reports unless explicitly opened through the Custom Work section.",
        "Work-log hours are recalculated on create and update from start/end times.",
        "Release-note files and project attachments each allow up to 20 MB per uploaded file.",
        "Planner items default to the project owner if no assignee is provided.",
        "Automation schedules are prevented from running more than once per configured period.",
    ]))
    lines.extend([""])

    lines.extend(section("8. Restrictions, Controls, and Observed Implementation Caveats", bullet_block(IMPLEMENTATION_CAVEATS)))
    lines.extend(section("9. Data Model Summary", [f"- `{table_name}`: {note}" for table_name, note in DATA_MODEL_NOTES.items()]))

    lines.extend(section("10. Route Inventory", []))
    lines.append("### Web Routes")
    lines.append("")
    for method, path, target, note in WEB_ROUTES:
        lines.append(f"- `{method} {path}` -> `{target}`: {note}")
    lines.append("")
    lines.append("### API Routes")
    lines.append("")
    for method, path, target, note in API_ROUTES:
        lines.append(f"- `{method} {path}` -> `{target}`: {note}")
    lines.append("")

    lines.extend(section("11. Deployment, Runtime, and Operations", bullet_block([
        "The local `.env.example` is now standardized on MySQL 8.",
        "Dockerfile installs `pdo_mysql` and builds the frontend with Node 20 before serving the Laravel app.",
        "Render deployment config expects MySQL environment variables rather than the old PostgreSQL mapping.",
        "The app uses Laravel's built-in `php artisan serve` in the container command, plus config/route/view cache generation, migrations, and seeding.",
        "Scheduled automation processing is implemented as the artisan command `automations:process` and should be wired into a scheduler or cron equivalent in production.",
        "Testing is configured for MySQL and requires a real MySQL test database rather than SQLite in-memory execution.",
    ])))

    lines.extend(section("12. Tech Stack Correction Work Completed", bullet_block([
        "Confirmed and preserved Laravel 13, Vue 3, Inertia.js, Tailwind CSS 4, and Vite as the live application stack.",
        "Removed PostgreSQL-specific runtime assumptions and connector customization from the application service provider.",
        "Changed environment defaults, queue defaults, test defaults, Docker runtime, and Render deployment config to MySQL 8.",
        "Replaced the placeholder Laravel README with project-specific stack and setup guidance.",
    ])))

    lines.extend(section("Appendix A. PHP Function Inventory", []))
    for rel_path, methods in collect_php_methods():
        lines.append(f"### {rel_path}")
        lines.append("")
        for method in methods:
            lines.append(f"- `{method}`")
        lines.append("")

    lines.extend(section("Appendix B. Frontend Page Inventory", []))
    lines.append("### Vue Pages")
    lines.append("")
    for page in collect_frontend_pages():
        lines.append(f"- `{page}`")
    lines.append("")
    lines.append("### Layouts and Shared Components")
    lines.append("")
    for item in collect_layouts_and_components():
        lines.append(f"- `{item}`")
    lines.append("")

    lines.extend(section("Appendix C. Enum Inventory", []))
    for enum_name, cases in collect_enum_values():
        lines.append(f"### {enum_name}")
        lines.append("")
        for case_name, value in cases:
            lines.append(f"- `{value}` ({case_name})")
        lines.append("")

    lines.extend(section("Appendix D. Schema Inventory Derived from Migrations", []))
    for table_name, columns in collect_schema_inventory().items():
        lines.append(f"### {table_name}")
        lines.append("")
        for column in columns:
            lines.append(f"- `{column}`")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def markdown_to_docx(markdown: str, output_path: Path) -> None:
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    for line in markdown.splitlines():
        if not line.strip():
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered_match:
            doc.add_paragraph(numbered_match.group(1), style="List Number")
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(line)

    doc.save(output_path)


def main() -> None:
    markdown = build_markdown()
    MARKDOWN_PATH.write_text(markdown, encoding="utf-8")
    markdown_to_docx(markdown, DOCX_PATH)
    print(f"Generated {MARKDOWN_PATH}")
    print(f"Generated {DOCX_PATH}")


if __name__ == "__main__":
    main()
